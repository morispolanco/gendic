import streamlit as st
import requests
import json
from docx import Document
from io import BytesIO

# Set page configuration
st.set_page_config(page_title="Generador de Diccionario por Campo de Estudio", page_icon="üìö", layout="wide")

# Function to create the information column
def crear_columna_info():
    st.markdown("""
    ## Sobre esta aplicaci√≥n

    Esta aplicaci√≥n es un Generador de Diccionario que permite a los usuarios crear un diccionario personalizado basado en un campo o √°rea de estudio espec√≠fico, incluyendo solo definiciones y referencias.

    ### C√≥mo usar la aplicaci√≥n:

    1. Ingrese un campo o √°rea de estudio de su inter√©s.
    2. Haga clic en "Generar t√©rminos" para obtener una lista de 101 t√©rminos relacionados.
    3. Edite la lista de t√©rminos seg√∫n sea necesario.
    4. Seleccione si desea generar definiciones para todos los t√©rminos o para un t√©rmino espec√≠fico.
    5. Haga clic en "Generar definiciones" para obtener las definiciones y referencias.
    6. Descargue un documento DOCX con las definiciones y referencias.

    ### Autor y actualizaci√≥n:
    **Moris Polanco**, [Fecha actual]

    ### C√≥mo citar esta aplicaci√≥n (formato APA):
    Polanco, M. (2024). *Generador de Diccionario por Campo de Estudio* [Aplicaci√≥n web]. [URL de la aplicaci√≥n]

    ---
    **Nota:** Esta aplicaci√≥n utiliza inteligencia artificial para generar t√©rminos y definiciones. Verifique la informaci√≥n con fuentes adicionales para un an√°lisis m√°s profundo.
    """)

# Titles and Main Column
st.title("Generador de Diccionario por Campo de Estudio")

col1, col2 = st.columns([1, 2])

with col1:
    crear_columna_info()

with col2:
    TOGETHER_API_KEY = st.secrets["TOGETHER_API_KEY"]
    SERPLY_API_KEY = st.secrets["SERPLY_API_KEY"]

    def generar_terminos(campo_estudio):
        url = "https://api.together.xyz/inference"
        payload = json.dumps({
            "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
            "prompt": f"Genera una lista de 101 t√©rminos relacionados con el campo de estudio: {campo_estudio}. Cada t√©rmino debe estar en una l√≠nea nueva.",
            "max_tokens": 2048,
            "temperature": 0.7,
            "top_p": 0.7,
            "top_k": 50,
            "repetition_penalty": 1,
        })
        headers = {
            'Authorization': f'Bearer {TOGETHER_API_KEY}',
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        terminos = response.json()['output']['choices'][0]['text'].strip().split('\n')
        return [termino.strip() for termino in terminos if termino.strip()]

    def buscar_informacion(query):
        url = f"https://api.serply.io/v1/scholar/q={query}"
        headers = {
            'X-Api-Key': SERPLY_API_KEY,
            'Content-Type': 'application/json',
            'X-Proxy-Location': 'US',
            'X-User-Agent': 'Mozilla/5.0'
        }
        response = requests.get(url, headers=headers)
        return response.json()

    def generar_definicion(termino, contexto):
        url = "https://api.together.xyz/inference"
        payload = json.dumps({
            "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
            "prompt": f"Contexto: {contexto}\n\nT√©rmino: {termino}\n\nProporciona una definici√≥n concisa y precisa del t√©rmino '{termino}' sin incluir conceptos relacionados, ant√≥nimos, sin√≥nimos u otra informaci√≥n adicional.\n\nDefinici√≥n:",
            "max_tokens": 1024,
            "temperature": 0.7,
            "top_p": 0.7,
            "top_k": 50,
            "repetition_penalty": 1,
            "stop": ["T√©rmino:"]
        })
        headers = {
            'Authorization': f'Bearer {TOGETHER_API_KEY}',
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        return response.json()['output']['choices'][0]['text'].strip()

    def create_docx(campo_estudio, terminos_definiciones, referencias):
        doc = Document()
        doc.add_heading(f'Diccionario de {campo_estudio}', 0)

        # Definiciones
        doc.add_heading('Definiciones', level=1)
        for termino, definicion in terminos_definiciones.items():
            doc.add_paragraph(f"{termino}: {definicion}")

        # Referencias
        doc.add_page_break()
        doc.add_heading('Referencias', level=1)
        for referencia in referencias:
            doc.add_paragraph(referencia, style='List Bullet')

        doc.add_paragraph('\nNota: Este documento fue generado por un asistente de IA. Verifique la informaci√≥n con fuentes acad√©micas para un an√°lisis m√°s profundo.')

        return doc

    def formatear_referencia_apa(ref):
        authors = ref.get('author', 'Autor desconocido')
        year = ref.get('year', 's.f.')
        title = ref.get('title', 'T√≠tulo desconocido')
        journal = ref.get('journal', '')
        volume = ref.get('volume', '')
        issue = ref.get('issue', '')
        pages = ref.get('pages', '')
        url = ref.get('url', '')

        reference = f"{authors} ({year}). {title}."
        if journal:
            reference += f" {journal}"
            if volume:
                reference += f", {volume}"
                if issue:
                    reference += f"({issue})"
            if pages:
                reference += f", {pages}"
        reference += f". {url}"
        
        return reference

    # Interfaz de usuario
    campo_estudio = st.text_input("Ingresa un campo o √°rea de estudio:")

    if st.button("Generar t√©rminos"):
        if campo_estudio:
            with st.spinner("Generando t√©rminos..."):
                terminos = generar_terminos(campo_estudio)
                st.session_state.terminos = terminos

    if 'terminos' in st.session_state:
        st.subheader("Lista de t√©rminos (editable):")
        terminos_editados = st.text_area("Edita los t√©rminos aqu√≠:", "\n".join(st.session_state.terminos), height=300)
        st.session_state.terminos_editados = terminos_editados.split('\n')

    if 'terminos_editados' in st.session_state:
        opcion_definicion = st.radio("Selecciona una opci√≥n:", ["Generar todas las definiciones", "Generar definici√≥n para un t√©rmino espec√≠fico"])

        if opcion_definicion == "Generar definici√≥n para un t√©rmino espec√≠fico":
            termino_seleccionado = st.selectbox("Selecciona un t√©rmino:", st.session_state.terminos_editados)

        if st.button("Generar definiciones"):
            with st.spinner("Generando definiciones y referencias..."):
                terminos_definiciones = {}
                todas_referencias = []

                if opcion_definicion == "Generar todas las definiciones":
                    for termino in st.session_state.terminos_editados:
                        resultados_busqueda = buscar_informacion(f"{termino} {campo_estudio}")
                        contexto = "\n".join([item["snippet"] for item in resultados_busqueda.get("results", [])])
                        definicion = generar_definicion(termino, contexto)
                        terminos_definiciones[termino] = definicion
                        referencias = [formatear_referencia_apa(item) for item in resultados_busqueda.get("results", [])]
                        todas_referencias.extend(referencias)
                else:
                    resultados_busqueda = buscar_informacion(f"{termino_seleccionado} {campo_estudio}")
                    contexto = "\n".join([item["snippet"] for item in resultados_busqueda.get("results", [])])
                    definicion = generar_definicion(termino_seleccionado, contexto)
                    terminos_definiciones[termino_seleccionado] = definicion
                    referencias = [formatear_referencia_apa(item) for item in resultados_busqueda.get("results", [])]
                    todas_referencias.extend(referencias)

                st.subheader("Definiciones generadas:")
                for termino, definicion in terminos_definiciones.items():
                    st.markdown(f"**{termino}**: {definicion}")

                st.subheader("Referencias:")
                for referencia in todas_referencias:
                    st.markdown(f"- {referencia}")

                # Bot√≥n para descargar el documento
                doc = create_docx(campo_estudio, terminos_definiciones, todas_referencias)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    label="Descargar diccionario en DOCX",
                    data=buffer,
                    file_name=f"Diccionario_{campo_estudio.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
